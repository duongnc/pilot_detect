from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def tao_khung_bao_cao():
    print("[INFO] Đang khởi tạo file Word...")
    
    # Tạo một tài liệu Word mới
    doc = Document()

    # Thêm Tiêu đề chính
    title = doc.add_heading('KHUNG BÁO CÁO PHẦN MỀM NHẬN DẠNG KHUÔN MẶT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Sinh viên thực hiện: [Tên của bạn]")
    doc.add_paragraph("Giảng viên hướng dẫn: [Tên GVHD]")
    doc.add_page_break() # Ngắt trang để bắt đầu vào nội dung

    # --- CHƯƠNG 1 ---
    doc.add_heading('CHƯƠNG 1. TỔNG QUAN', level=1)
    doc.add_heading('1.1. Giới thiệu về kỹ thuật nhận dạng khuôn mặt và phân tích biểu cảm', level=2)
    doc.add_heading('1.1.1. Nhận dạng khuôn mặt trong đời sống và hệ thống giám sát (GCS)', level=3)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('1.1.2. Hệ thống giám sát có ứng dụng trí tuệ nhân tạo (AI)', level=3)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('1.2. Tổng quan bài toán phát hiện và nhận dạng trạng thái khuôn mặt', level=2)
    doc.add_heading('1.2.1. Giới thiệu bài toán', level=3)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('1.2.2. Hướng tiếp cận bài toán bằng mô hình học sâu (Deep Learning)', level=3)
    doc.add_paragraph('Nội dung trình bày ở đây...')

    # --- CHƯƠNG 2 ---
    doc.add_heading('CHƯƠNG 2. CƠ SỞ LÝ THUYẾT', level=1)
    doc.add_heading('2.1. Kiến trúc tổng thể của hệ thống nhận dạng đa phương thức (Multi-modal)', level=2)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('2.2. Mạng nơ-ron phát hiện và cắt tỉa khuôn mặt (MTCNN / Face Detector)', level=2)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('2.3. Rút trích đặc trưng khuôn mặt diện rộng (Mạng DAN / ResNet18 Backbone)', level=2)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('2.4. Rút trích đặc trưng cử động cơ mặt cục bộ (Mạng VGG19 Backbone)', level=2)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('2.5. Cơ sở lý thuyết về Dung hợp dữ liệu (Multi-Modal Fusion & Adaptive Weights)', level=2)
    doc.add_paragraph('Nội dung trình bày ở đây...')

    # --- CHƯƠNG 3 ---
    doc.add_heading('CHƯƠNG 3. TRIỂN KHAI ĐÁNH GIÁ', level=1)
    doc.add_heading('3.1. Thiết lập môi trường và thư viện', level=2)
    doc.add_heading('3.1.1. Cài đặt (PyTorch, OpenCV, Môi trường ảo venv, PyInstaller)', level=3)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('3.2. Xây dựng đường ống xử lý thời gian thực (Real-time Pipeline)', level=2)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('3.3. Tối ưu hóa trên các nền tảng phần cứng (CUDA, Apple MPS, CPU)', level=2)
    doc.add_paragraph('Nội dung trình bày ở đây...')

    # --- CHƯƠNG 4 ---
    doc.add_heading('CHƯƠNG 4. KẾT QUẢ', level=1)
    doc.add_heading('4.1. Dữ liệu (Tập dữ liệu RAF-DB và Tập dữ liệu 7 cảm xúc tự thu thập)', level=2)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('4.2. Độ đo đánh giá', level=2)
    doc.add_heading('4.2.1. Ma trận nhầm lẫn (Confusion Matrix)', level=3)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('4.2.2. Intersection over Union (IoU)', level=3)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('4.2.3. Receiver operating characteristic curve (ROC Curve)', level=3)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('4.2.4. Độ chính xác trung bình (mAP / Accuracy / F1-Score)', level=3)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('4.3. Kết quả thực hiện', level=2)
    doc.add_paragraph('Nội dung trình bày ở đây...')
    
    doc.add_heading('4.4. Kết luận và Hướng phát triển', level=2)
    doc.add_paragraph('Nội dung trình bày ở đây...')

    # --- TÀI LIỆU THAM KHẢO ---
    doc.add_heading('TÀI LIỆU THAM KHẢO', level=1)
    doc.add_paragraph('[1] Tên bài báo/tài liệu tham khảo...')

    # Lưu lại file Word
    file_name = "Khung_Bao_Cao_Nhan_Dang_Khuon_Mat.docx"
    doc.save(file_name)
    print(f"[THÀNH CÔNG] Đã tạo xong file: {file_name}")

if __name__ == "__main__":
    tao_khung_bao_cao()